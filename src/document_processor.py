"""
Functions for processing PDF documents and other file formats.
Handles direct PDF submission for AI models.
"""

import os
import base64
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple, Union
import time
import random

from PIL import Image
try:
    import docx
except ImportError:
    print("Warning: python-docx is not installed. DOCX text extraction will be unavailable.")
    docx = None

import config

def is_direct_pdf_supported(model_type: str) -> bool:
    """
    Check if the specified model type supports direct PDF input.
    
    Args:
        model_type: Type of AI model (e.g., 'anthropic', 'openai')
        
    Returns:
        True if direct PDF input is supported, False otherwise
    """
    # Currently only Anthropic models support direct PDF input
    return model_type.lower() == 'anthropic'

def compress_image(image: Image.Image, quality: int = None, max_size: Tuple[int, int] = None) -> Image.Image:
    """
    Compress and resize an image to reduce file size.
    
    Args:
        image: PIL Image object
        quality: JPEG quality (1-100)
        max_size: Maximum width and height
        
    Returns:
        Compressed PIL Image object
    """
    if quality is None:
        quality = config.IMAGE_SETTINGS['quality']
    
    if max_size is None:
        max_size = config.IMAGE_SETTINGS['max_size']
        
    # Create a copy to avoid modifying the original
    image_copy = image.copy()
    
    # Resize if needed
    if image_copy.width > max_size[0] or image_copy.height > max_size[1]:
        image_copy.thumbnail(max_size, Image.LANCZOS)
    
    return image_copy

def image_to_base64(image: Image.Image, format: str = None, quality: int = None) -> str:
    """
    Convert PIL Image to base64 string with compression.
    
    Args:
        image: PIL Image object
        format: Output format (e.g., 'JPEG', 'PNG')
        quality: Image quality (1-100)
        
    Returns:
        Base64-encoded image string
    """
    if format is None:
        format = config.IMAGE_SETTINGS['format']
        
    if quality is None:
        quality = config.IMAGE_SETTINGS['quality']
        
    buffered = BytesIO()
    image.save(buffered, format=format, quality=quality, optimize=True)
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def file_to_base64(file_path: str) -> str:
    """
    Convert any file directly to base64 string.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Base64-encoded string
    """
    try:
        with open(file_path, 'rb') as file:
            encoded_string = base64.b64encode(file.read()).decode('utf-8')
        return encoded_string
    except Exception as e:
        print(f"Error converting file to base64: {str(e)}")
        return ""

def get_file_media_type(file_ext: str) -> str:
    """
    Get the media type for a file extension.
    
    Args:
        file_ext: File extension (with dot, e.g., '.pdf')
        
    Returns:
        Media type string
    """
    # Use the media types from config
    return config.FILE_TYPES["media_types"].get(file_ext.lower(), "application/octet-stream")

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
    """
    if docx is None:
        return "Cannot extract text from DOCX: python-docx not installed"
        
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def process_submission_file(file_path: str, model_type: str) -> Dict[str, Any]:
    """
    Process a submission file based on its type and the target model.
    
    Args:
        file_path: Path to the submission file
        model_type: Type of model that will process the file
        
    Returns:
        Dictionary with processed file data:
        - type: 'file_base64', 'image_base64', or 'error'
        - content: base64 string content
        - media_type: MIME type of the file
        - error: Error message if processing failed
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        # For images, compress and convert to base64
        if file_ext in config.FILE_TYPES["image_extensions"]:
            try:
                img = Image.open(file_path)
                img = compress_image(img)
                image_data = image_to_base64(img)
                return {
                    "type": "image_base64",
                    "content": image_data,
                    "media_type": get_file_media_type(file_ext)
                }
            except Exception as e:
                return {
                    "type": "error",
                    "error": f"Error processing image file: {str(e)}"
                }
                
        # For PDF documents
        elif file_ext == '.pdf':
            file_data = file_to_base64(file_path)
            return {
                "type": "file_base64",
                "content": file_data,
                "media_type": "application/pdf"
            }
            
        # For non-PDF documents
        elif file_ext in config.FILE_TYPES["document_extensions"]:
            # If using Anthropic, which only accepts PDF documents
            if model_type.lower() == 'anthropic':
                # For documents that Anthropic can't process directly,
                # we'll convert them to text and send as text content
                try:
                    # Extract text from DOCX files
                    if file_ext == '.docx':
                        content = extract_text_from_docx(file_path)
                        return {
                            "type": "text",
                            "content": f"[Content from {os.path.basename(file_path)}]:\n\n{content}"
                        }
                    # Simple text extraction for plain text files
                    elif file_ext in ['.txt', '.md', '.csv']:
                        with open(file_path, 'r', errors='ignore') as f:
                            content = f.read()
                        return {
                            "type": "text",
                            "content": f"[Content from {os.path.basename(file_path)}]:\n\n{content}"
                        }
                    # For other document types we'll handle as base64 image
                    # Anthropic will process it as an image rather than a document
                    else:
                        file_data = file_to_base64(file_path)
                        return {
                            "type": "image_base64",
                            "content": file_data,
                            "media_type": "application/octet-stream",
                            "filename": os.path.basename(file_path)
                        }
                except Exception as e:
                    return {
                        "type": "error",
                        "error": f"Error processing non-PDF document for Anthropic: {str(e)}"
                    }
            # For other models that can handle different document types
            else:
                file_data = file_to_base64(file_path)
                media_type = get_file_media_type(file_ext)
                return {
                    "type": "file_base64",
                    "content": file_data,
                    "media_type": media_type
                }
        else:
            return {
                "type": "error",
                "error": f"Unsupported file type: {file_ext}"
            }
            
    except Exception as e:
        return {
            "type": "error",
            "error": f"Error processing file: {str(e)}"
        }

def prepare_submission_for_model(submission_files: List[Dict[str, Any]], model_type: str) -> Dict[str, Any]:
    """
    Process submission files for model consumption.
    
    Args:
        submission_files: List of submission file dictionaries
        model_type: Type of model to use (anthropic, openai)
        
    Returns:
        Dictionary with processed content
    """
    try:
        content = []
        total_size = 0
        size_limit = 20 * 1024 * 1024  # 20MB limit for total submission size
        compressed_size_limit = 6 * 1024 * 1024  # 6MB limit after compression
        
        for file_info in submission_files:
            file_path = file_info.get("path")
            if not os.path.exists(file_path):
                continue
                
            file_ext = os.path.splitext(file_path.lower())[1]
            file_size = os.path.getsize(file_path)
            total_size += file_size
            
            # Process based on file type
            if file_ext in config.FILE_TYPES["document_extensions"]:
                # Document file (PDF, DOC, etc.)
                try:
                    is_large_file = file_size > 5 * 1024 * 1024  # 5MB threshold
                    
                    # For large PDFs, use compression
                    if file_ext == '.pdf' and is_large_file:
                        compressed_path = compress_pdf(file_path)
                        if compressed_path:
                            file_path = compressed_path
                    
                    # Convert to base64
                    b64_data = file_to_base64(file_path)
                    
                    # If still too large after compression, skip with warning
                    if len(b64_data) > compressed_size_limit:
                        print(f"Warning: File {os.path.basename(file_path)} is too large even after compression. Using text extraction instead.")
                        # Fall back to text extraction
                        extracted_text = extract_text_from_document(file_path)
                        content.append({
                            "type": "text",
                            "filename": os.path.basename(file_path),
                            "content": f"EXTRACTED TEXT FROM {os.path.basename(file_path)}:\n\n{extracted_text}"
                        })
                    else:
                        content.append({
                            "type": "file_base64",
                            "filename": os.path.basename(file_path),
                            "content": b64_data,
                            "media_type": get_file_media_type(file_ext)
                        })
                        
                except Exception as e:
                    # If document processing fails, try to extract text
                    print(f"Error processing document {file_path}: {str(e)}. Attempting text extraction.")
                    try:
                        extracted_text = extract_text_from_document(file_path)
                        content.append({
                            "type": "text",
                            "filename": os.path.basename(file_path),
                            "content": f"EXTRACTED TEXT FROM {os.path.basename(file_path)}:\n\n{extracted_text}"
                        })
                    except Exception as text_error:
                        print(f"Failed to extract text from {file_path}: {str(text_error)}")
            
            elif file_ext in config.FILE_TYPES["image_extensions"]:
                # Image file
                try:
                    is_large_image = file_size > 1 * 1024 * 1024  # 1MB threshold
                    
                    # For large images, compress them
                    if is_large_image:
                        compressed_path = compress_image(file_path)
                        if compressed_path:
                            file_path = compressed_path
                    
                    # Convert to base64
                    b64_data = file_to_base64(file_path)
                    
                    # If still too large, skip
                    if len(b64_data) > compressed_size_limit:
                        print(f"Warning: Image {os.path.basename(file_path)} is too large even after compression. Skipping.")
                        continue
                        
                    content.append({
                        "type": "image_base64",
                        "filename": os.path.basename(file_path),
                        "content": b64_data,
                        "media_type": get_file_media_type(file_ext)
                    })
                except Exception as e:
                    print(f"Error processing image {file_path}: {str(e)}")
                    
            elif file_ext == '.txt':
                # Text file - read contents directly
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        file_text = f.read()
                        
                    content.append({
                        "type": "text",
                        "filename": os.path.basename(file_path),
                        "content": file_text
                    })
                except Exception as e:
                    print(f"Error reading text file {file_path}: {str(e)}")
        
        # Check if total size is too large and apply more aggressive measures if needed
        if total_size > size_limit:
            print(f"Warning: Total submission size ({total_size/1024/1024:.2f}MB) exceeds recommended limit. Using text extraction for all documents.")
            # Reset content and try with text extraction for all files
            content = []
            for file_info in submission_files:
                file_path = file_info.get("path")
                if not os.path.exists(file_path):
                    continue
                    
                file_ext = os.path.splitext(file_path.lower())[1]
                
                if file_ext in config.FILE_TYPES["document_extensions"]:
                    try:
                        extracted_text = extract_text_from_document(file_path)
                        content.append({
                            "type": "text",
                            "filename": os.path.basename(file_path),
                            "content": f"EXTRACTED TEXT FROM {os.path.basename(file_path)}:\n\n{extracted_text}"
                        })
                    except Exception as e:
                        print(f"Failed to extract text from {file_path}: {str(e)}")
                elif file_ext == '.txt':
                    # Text file - read contents directly
                    try:
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            file_text = f.read()
                            
                        content.append({
                            "type": "text",
                            "filename": os.path.basename(file_path),
                            "content": file_text
                        })
                    except Exception as e:
                        print(f"Error reading text file {file_path}: {str(e)}")
                        
        # If no content was processed successfully, return error
        if not content:
            return {
                "type": "error",
                "error": "No files were processed successfully"
            }
            
        return {
            "type": "success",
            "content": content
        }
        
    except Exception as e:
        return {
            "type": "error",
            "error": f"Error processing submission: {str(e)}"
        }

def compress_pdf(pdf_path: str) -> Optional[str]:
    """
    Compress a PDF file to reduce its size.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Path to the compressed PDF or None if compression failed
    """
    try:
        # Import necessary libraries
        from PyPDF2 import PdfReader, PdfWriter
        import tempfile
        
        # Create a temporary file to store the compressed PDF
        fd, compressed_path = tempfile.mkstemp(suffix='.pdf')
        os.close(fd)
        
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        # Add each page with compression enabled
        for page in reader.pages:
            writer.add_page(page)
            
        # Use more aggressive compression
        writer.add_metadata(reader.metadata)
        
        # Write the compressed file with compression enabled
        with open(compressed_path, 'wb') as output_file:
            writer.write(output_file)
            
        # Return the path to the compressed file
        return compressed_path
    except Exception as e:
        print(f"Error compressing PDF {pdf_path}: {str(e)}")
        return None

def compress_image(image_path: str) -> Optional[str]:
    """
    Compress an image file to reduce its size.
    
    Args:
        image_path: Path to the image file
        
    Returns:
        Path to the compressed image or None if compression failed
    """
    try:
        # Import necessary libraries
        from PIL import Image
        import tempfile
        
        # Create a temporary file to store the compressed image
        fd, compressed_path = tempfile.mkstemp(suffix=os.path.splitext(image_path)[1])
        os.close(fd)
        
        # Open the image
        img = Image.open(image_path)
        
        # Resize large images
        max_size = config.IMAGE_SETTINGS.get("max_size", (800, 800))
        if img.width > max_size[0] or img.height > max_size[1]:
            img.thumbnail(max_size, Image.LANCZOS)
            
        # Save with compression
        img.save(
            compressed_path, 
            format=config.IMAGE_SETTINGS.get("format", "JPEG"),
            optimize=True,
            quality=config.IMAGE_SETTINGS.get("quality", 40)
        )
        
        # Return the path to the compressed file
        return compressed_path
    except Exception as e:
        print(f"Error compressing image {image_path}: {str(e)}")
        return None

def extract_text_from_document(doc_path: str) -> str:
    """
    Extract text from a document file.
    
    Args:
        doc_path: Path to the document file
        
    Returns:
        Extracted text
    """
    try:
        file_ext = os.path.splitext(doc_path.lower())[1]
        
        # Handle PDF files
        if file_ext == '.pdf':
            try:
                import pypdf
                extracted_text = ""
                
                with open(doc_path, 'rb') as file:
                    pdf = pypdf.PdfReader(file)
                    # Extract text from each page
                    for page_num in range(len(pdf.pages)):
                        page = pdf.pages[page_num]
                        extracted_text += f"--- Page {page_num + 1} ---\n"
                        extracted_text += page.extract_text() or "[No extractable text on this page]"
                        extracted_text += "\n\n"
                        
                return extracted_text
            except Exception as e:
                print(f"Error extracting text from PDF: {str(e)}")
                return f"[Failed to extract text from PDF: {str(e)}]"
                
        # Handle DOCX files
        elif file_ext == '.docx':
            try:
                import docx
                doc = docx.Document(doc_path)
                return "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                print(f"Error extracting text from DOCX: {str(e)}")
                return f"[Failed to extract text from DOCX: {str(e)}]"
                
        # Handle other document types
        else:
            return f"[Text extraction not supported for {file_ext} files]"
            
    except Exception as e:
        print(f"Error extracting text from document: {str(e)}")
        return f"[Text extraction failed: {str(e)}]"

def get_allowed_file_extensions():
    """
    Get a list of all allowed file extensions.
    
    Returns:
        List of allowed file extensions (e.g., ['.pdf', '.jpg'])
    """
    # Combine document and image extensions from config
    return config.FILE_TYPES["document_extensions"] + config.FILE_TYPES["image_extensions"]

def analyze_lab_questions_with_ai(pdf_path: str, model_type: str = "anthropic") -> Dict[str, Any]:
    """
    Analyze a lab question PDF using AI to automatically determine structure and grading criteria.
    
    Args:
        pdf_path: Path to the PDF file containing lab questions
        model_type: Type of AI model to use for analysis
        
    Returns:
        Dictionary with lab structure information:
        - name: Lab name/title
        - total_questions: Number of questions/sections
        - question_names: List of question/section names
        - marks_per_question: Points per question (if evenly distributed)
        - total_marks: Total available marks
    """
    # Check if the specified model supports direct PDF input
    if not is_direct_pdf_supported(model_type):
        return {"error": True, "error_message": f"Model {model_type} doesn't support direct PDF analysis"}
    
    # Import model class based on model_type
    if model_type.lower() == "anthropic":
        from model_interface import AnthropicModel
        model = AnthropicModel()
    else:
        return {"error": True, "error_message": f"Unsupported model type: {model_type}"}
    
    # Convert PDF to base64 for AI processing
    pdf_content = file_to_base64(pdf_path)
    
    # Create the analysis prompt
    system_prompt = """You are a teaching assistant analyzing lab instruction documents. 
Your task is to extract the structure of the lab questions, including titles, point values, and number of questions.
Return your analysis as structured JSON only, with no additional explanation or text."""
    
    message_content = [
        {"type": "text", "text": "Please analyze this lab instruction document and extract the following information:\n\n"
                                 "1. The lab title/name\n"
                                 "2. The total number of questions or sections to be completed\n"
                                 "3. The name of each question or section\n"
                                 "4. The points allocated to each question (if specified)\n"
                                 "5. The total points for the lab\n\n"
                                 "Return the analysis as a JSON object with this structure:\n"
                                 "{\n"
                                 "  \"name\": \"Lab Title\",\n"
                                 "  \"total_questions\": 4,\n"
                                 "  \"question_names\": [\"Question 1 Title\", \"Question 2 Title\", ...],\n"
                                 "  \"question_points\": [25, 25, ...],\n"
                                 "  \"total_marks\": 100\n"
                                 "}\n\n"
                                 "If points are not explicitly specified, distribute 100 points equally among the questions."
        },
        {"type": "document", 
         "source": {
             "type": "base64",
             "media_type": "application/pdf",
             "data": pdf_content
         }}
    ]
    
    # Send to AI for analysis
    try:
        result = model.send_message(message_content, system_prompt)
        
        # Ensure we have the required fields
        if "name" in result and "total_questions" in result and "question_names" in result:
            # Calculate marks_per_question if not provided
            if "question_points" in result:
                # Use provided point values
                points = result["question_points"]
                result["total_marks"] = sum(points)
            else:
                # Equal distribution of points
                if "total_marks" not in result:
                    result["total_marks"] = 100
                result["marks_per_question"] = result["total_marks"] // result["total_questions"]
                result["question_points"] = [result["marks_per_question"]] * result["total_questions"]
            
            # Add file path for reference
            result["pdf_path"] = pdf_path
            return result
        else:
            return {
                "error": True, 
                "error_message": "AI analysis did not return expected structure",
                "ai_response": result
            }
            
    except Exception as e:
        return {"error": True, "error_message": f"Error during AI analysis: {str(e)}"} 